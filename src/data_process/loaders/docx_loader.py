# -*- coding: utf-8 -*-
"""
Word (.docx) 文档加载器

支持文本和表格提取，以及嵌入图片的 OCR 识别。
"""

import logging
from io import BytesIO
from typing import List, Optional

import numpy as np
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image
from langchain_core.documents import Document

from src.data_process.loaders.base import BaseLoader

logger = logging.getLogger(__name__)


class DocxLoader(BaseLoader):
    """Word 文档加载器

    支持文本、表格提取和嵌入图片 OCR。
    """

    def __init__(
        self,
        file_path: str,
        encoding: Optional[str] = None,
        ocr_enabled: bool = True,
        use_cuda: bool = True,
        **kwargs,
    ) -> None:
        """初始化 Docx 加载器

        Args:
            file_path: Word 文件路径
            encoding: 编码（docx 通常为二进制，无需指定）
            ocr_enabled: 是否启用图片 OCR
            use_cuda: OCR 是否使用 GPU
            **kwargs: 其他参数
        """
        super().__init__(file_path, encoding, **kwargs)
        self.ocr_enabled = ocr_enabled
        self.use_cuda = use_cuda
        self._ocr_engine = None

    @property
    def ocr_engine(self):
        """延迟初始化 OCR 引擎"""
        if self._ocr_engine is None:
            try:
                from rapidocr_paddle import RapidOCR
                self._ocr_engine = RapidOCR(
                    det_use_cuda=self.use_cuda,
                    cls_use_cuda=self.use_cuda,
                    rec_use_cuda=self.use_cuda,
                )
            except ImportError:
                from rapidocr_onnxruntime import RapidOCR
                self._ocr_engine = RapidOCR()
        return self._ocr_engine

    def load(self) -> List[Document]:
        """加载 Word 文档

        Returns:
            List[Document]: 文档列表
        """
        from docx.document import Document as DocxDocType

        doc = DocxDocument(str(self.file_path))
        text_parts = []

        for block in self._iter_block_items(doc):
            if isinstance(block, Paragraph):
                text_parts.append(block.text.strip())
                if self.ocr_enabled:
                    ocr_text = self._extract_paragraph_images(block)
                    if ocr_text:
                        text_parts.append(ocr_text)
            elif isinstance(block, Table):
                table_text = self._extract_table_text(block)
                text_parts.append(table_text)

        full_text = "\n\n".join(filter(None, text_parts))

        metadata = {
            "source": str(self.file_path),
            "file_name": self.file_path.name,
            "file_type": "docx",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }

        return [Document(page_content=full_text, metadata=metadata)]

    def _iter_block_items(self, parent: "DocxDocType | _Cell"):
        """遍历文档块（段落和表格）"""
        from docx.document import Document as DocxDocType

        if isinstance(parent, DocxDocType):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Unsupported parent type")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _extract_paragraph_images(self, paragraph: Paragraph) -> str:
        """提取段落中的嵌入图片并进行 OCR

        Args:
            paragraph: 段落对象

        Returns:
            str: OCR 识别文本
        """
        from docx.parts.image import ImagePart

        try:
            images = paragraph._element.xpath(".//pic:pic")
        except Exception:
            return ""

        if not images:
            return ""

        doc = paragraph._element.getparent().getparent()
        while doc.tag != "w:document":
            doc = doc.getparent()

        resp_parts = []

        for image in images:
            try:
                for img_id in image.xpath(".//a:blip/@r:embed"):
                    part = doc.part.related_parts.get(img_id)
                    if isinstance(part, ImagePart):
                        img = Image.open(BytesIO(part._blob))
                        img_array = np.array(img)
                        result, _ = self.ocr_engine(img_array)
                        if result:
                            ocr_lines = [line[1] for line in result]
                            resp_parts.extend(ocr_lines)
            except Exception as e:
                logger.debug(f"段落图片 OCR 处理失败: {e}")
                continue

        return "\n".join(resp_parts)

    def _extract_table_text(self, table: Table) -> str:
        """提取表格文本

        Args:
            table: 表格对象

        Returns:
            str: 表格文本
        """
        rows_text = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs)
                cells.append(cell_text)
            rows_text.append(" | ".join(cells))
        return "\n".join(rows_text)
